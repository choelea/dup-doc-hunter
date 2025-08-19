import os
import uuid
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
import boto3

# ===== MinIO 配置 =====
MINIO_ENDPOINT = "http://10.3.70.127:9000"
MINIO_ACCESS_KEY = "minioadmin"
MINIO_SECRET_KEY = "minioadmin"
MINIO_BUCKET = "md-images"

s3 = boto3.client(
    "s3",
    endpoint_url=MINIO_ENDPOINT,
    aws_access_key_id=MINIO_ACCESS_KEY,
    aws_secret_access_key=MINIO_SECRET_KEY,
    region_name="us-east-1"
)

def upload_to_minio(local_file, bucket, object_name=None):
    if object_name is None:
        object_name = os.path.basename(local_file)
    s3.upload_file(local_file, bucket, object_name)
    return f"{MINIO_ENDPOINT}/{bucket}/{object_name}"

def escape_pipe(text):
    return text.replace("|", "\\|")

def table_to_markdown(table):
    rows = []
    for i, row in enumerate(table.rows):
        cells = [escape_pipe(cell.text.strip()) for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            sep = ["---"] * len(row.cells)
            rows.append("| " + " | ".join(sep) + " |")
    return "\n".join(rows) + "\n"

def extract_inline_image(run, output_dir="output/media"):
    """提取 run 内图片并上传 MinIO，返回 Markdown URL"""
    os.makedirs(output_dir, exist_ok=True)
    # 遍历 run 元素子孙节点，找到 pic:pic 元素
    for child in run._element.iter():
        if child.tag.endswith('pic'):
            # 获取 blip 节点
            blip = child.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
            if blip is not None:
                rId = blip.get(qn("r:embed"))
                image_part = run.part.related_parts[rId]
                ext = image_part.content_type.split("/")[-1]
                img_name = f"{uuid.uuid4().hex}.{ext}"
                local_path = os.path.join(output_dir, img_name)
                with open(local_path, "wb") as f:
                    f.write(image_part.blob)
                # 上传到 MinIO
                url = upload_to_minio(local_path, MINIO_BUCKET, f"media/{img_name}")
                # 删除本地临时文件
                os.remove(local_path)
                return url
    return None

def convert_docx_to_md(docx_path, output_dir="output"):
    os.makedirs(output_dir, exist_ok=True)
    md_path = os.path.join(output_dir, f"{os.path.splitext(os.path.basename(docx_path))[0]}.md")

    doc = Document(docx_path)
    md_lines = []

    for block in doc.element.body:
        tag = block.tag.split("}")[1]

        # 表格
        if tag == "tbl":
            table = next((t for t in doc.tables if t._element == block), None)
            if table:
                md_lines.append(table_to_markdown(table))
            continue

        # 段落
        if tag == "p":
            para = next((p for p in doc.paragraphs if p._element == block), None)
            if para:
                style = para.style.name.lower() if para.style else "normal"
                line_parts = []

                for run in para.runs:
                    text = run.text
                    img_url = extract_inline_image(run, os.path.join(output_dir, "media"))
                    if img_url:
                        line_parts.append(f"![image]({img_url})")
                    if text:
                        line_parts.append(text)
                line = "".join(line_parts).strip()
                if not line:
                    continue

                # 标题
                if "heading" in style:
                    level = int(style.replace("heading ", ""))
                    md_lines.append("#" * level + " " + line + "\n")
                # 列表
                elif style.startswith("list"):
                    md_lines.append(f"- {line}\n")
                else:
                    md_lines.append(line + "\n")

    # 保存 Markdown
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))

    return md_path

if __name__ == "__main__":
    docx_file = "input/企业AI 知识文档治理规范性建议.docx"
    md_file = convert_docx_to_md(docx_file)
    print(f"✅ 转换完成，Markdown 文件保存在: {md_file}")